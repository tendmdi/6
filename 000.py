import tkinter as tk
from module import mass
from module import planka
from module import svet
from docx import Document
from openpyxl import Workbook



def fu(particle):
    global selected_particle
    selected_particle = particle

def udele():
    global selected_particle
    if selected_particle == 'Протон':
        z = mass.zp
        m = mass.mp
    elif selected_particle == 'Электрон':
        z = mass.ze
        m = mass.me
    elif selected_particle == 'Нейтрон':
        z = 0
        m = mass.mn
    else:
        return
    
    k = z / m
    result_label.config(text=f"Удельный заряд: {k}")

    volna = (planka.h) / (svet.c * m)

    doc = Document()
    doc.add_paragraph(f'Удельный заряд частицы "{selected_particle}": {k}')
    doc.add_paragraph(f'Комптоновская длина волны частицы "{selected_particle}": {volna}')
    doc.save('вордик.docx')

    wb = Workbook()
    ws = wb.active
    ws.append(['Удельный заряд', k])
    ws.append(['Комптоновская длина волны', volna])
    wb.save('экселька.xlsx')

root = tk.Tk()

canvas = tk.Canvas(root, width=900, height=400)
canvas.pack()

r = 50
electron = canvas.create_oval(100-r, 200-r, 100+r, 200+r, fill='pink')  
canvas.create_text(100, 200, text="-", font=("Arial", 30), fill="black")

proton = canvas.create_oval(800-r, 200-r, 800+r, 200+r, fill='light blue')  
canvas.create_text(800, 200, text="+", font=("Arial", 30), fill="black")

neutron = canvas.create_oval(450-r, 200-r, 450+r, 200+r, fill='purple')  
canvas.create_text(450, 200, text="", font=("Arial", 30), fill="black")

root.title("Расчет удельного заряда")

z_label = tk.Label(root, text="Заряд частицы:")
z_label.pack()

z_entry = tk.Entry(root)
z_entry.pack()

m_label = tk.Label(root, text="Масса частицы:")
m_label.pack()

m_entry = tk.Entry(root)
m_entry.pack()

calculate_button = tk.Button(root, text="Рассчитать", command=udele)
calculate_button.pack()

result_label = tk.Label(root, text="")
result_label.pack()

particle_selection_frame = tk.Frame(root)
particle_selection_frame.pack()

particle_selection_label = tk.Label(particle_selection_frame, text="Выберите частицу:")
particle_selection_label.pack()

particles = ['Протон', 'Электрон', 'Нейтрон']

for particle in particles:
    particle_button = tk.Radiobutton(particle_selection_frame, text=particle, value=particle, command=lambda p=particle: fu(p))
    particle_button.pack()

root.mainloop()
